
from docx import Document
import os

def export_to_all(filename: str, content: str, tags: list):
    doc = Document()
    doc.add_heading("Harmonic Resonance Master – Grant Proposal", level=1)
    doc.add_paragraph(content)
    export_path = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    doc.save(export_path)
